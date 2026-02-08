import streamlit as st
import os
import tempfile
import psycopg2
from dotenv import load_dotenv
from psycopg2.extras import execute_values
from datetime import datetime
import json
from typing import List, Optional
import requests

# Импорты LlamaIndex
from llama_index.core import (
    VectorStoreIndex,
    Document,
    StorageContext,
    Settings,
    load_index_from_storage,
)
from llama_index.llms.openrouter import OpenRouter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.postgres import PGVectorStore
from llama_index.core.node_parser import SemanticSplitterNodeParser
from llama_index.core.ingestion import IngestionPipeline

# Импорты для работы с документами
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import hashlib

load_dotenv()

# Настройки страницы
st.set_page_config(
    page_title="спинтехтёнок",
    page_icon="🎓",
    layout="wide"
)

# Конфигурация
class Config:
    # Настройки PostgreSQL
    DB_NAME = os.environ.get("DB_NAME")
    DB_USER = os.environ.get("DB_LOGIN")
    DB_PASSWORD = os.environ.get("DB_PASSWORD")
    DB_HOST = "localhost"
    DB_PORT = int("5432")
    
    # Настройки OpenRouter
    OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
    OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
    
    # Настройки модели
    LLM_MODEL = os.environ.get("OPENROUTER_MODEL")
    EMBEDDING_MODEL = os.environ.get("EMBEDDING_MODEL")
    EMBEDDING_DIM = int(os.environ.get("EMBEDDING_DIM"))
    
    # Размер чанков
    CHUNK_SIZE = 512
    CHUNK_OVERLAP = 50

def init_db():
    try:
        conn = psycopg2.connect(
            dbname=Config.DB_NAME,
            user=Config.DB_USER,
            password=Config.DB_PASSWORD,
            host=Config.DB_HOST,
            port=Config.DB_PORT
        )
        return conn
    except Exception as e:
        st.error(f"Ошибка подключения к базе данных: {e}")
        return None

def create_tables():
    conn = init_db()
    if not conn:
        return
    
    cur = conn.cursor()
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id SERIAL PRIMARY KEY,
            filename VARCHAR(255) NOT NULL,
            file_type VARCHAR(50) NOT NULL,
            file_size BIGINT,
            file_hash VARCHAR(64) UNIQUE,
            upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            content TEXT,
            metadata JSONB,
            processed BOOLEAN DEFAULT FALSE
        )
    """)
    
    conn.commit()
    cur.close()
    conn.close()

    # Таблица для чанков будет создана LlamaIndex

@st.cache_resource
def init_models():
    llm = OpenRouter(
        model=Config.LLM_MODEL,
        api_key=Config.OPENROUTER_API_KEY,
        base_url=Config.OPENROUTER_BASE_URL,
        temperature=0.1,
        context_window=4096
    )
    
    embed_model = HuggingFaceEmbedding(
        model_name=Config.EMBEDDING_MODEL
    )
    
    Settings.llm = llm
    Settings.embed_model = embed_model
    
    return llm, embed_model

@st.cache_resource
def init_vector_store():
    try:
        vector_store = PGVectorStore.from_params(
            database=Config.DB_NAME,
            host=Config.DB_HOST,
            password=Config.DB_PASSWORD,
            port=Config.DB_PORT,
            user=Config.DB_USER,
            table_name="document_chunks",
            embed_dim=Config.EMBEDDING_DIM
        )
        return vector_store
    except Exception as e:
        st.error(f"Ошибка инициализации векторного хранилища: {e}")
        return None

def extract_text_from_file(file_path, file_type):
    text = ""
    
    try:
        if file_type == "pdf":
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        
        elif file_type == "docx":
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        
        elif file_type == "txt":
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
        
        return text.strip()
    
    except Exception as e:
        st.error(f"Ошибка чтения файла: {e}")
        return ""

from bs4 import BeautifulSoup

def extract_text_from_url(url):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Определяем кодировку
        response.encoding = response.apparent_encoding
        
        # Используем BeautifulSoup для извлечения текста
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Удаляем ненужные элементы
        for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 
                            'iframe', 'noscript', 'svg', 'form', 'button', 'input']):
            element.decompose()
        
        # Получаем текст из основных контентных элементов
        text_parts = []
        
        # Приоритет: основные контентные теги
        content_tags = ['article', 'main', 'section.content', 'div.content', 
                       'div.article', 'div.post', 'div.entry-content']
        
        for tag in content_tags:
            if ',' in tag:
                selector, class_name = tag.split('.')
                elements = soup.find_all(selector, class_=class_name)
            else:
                elements = soup.find_all(tag)
            
            if elements:
                for element in elements:
                    text = element.get_text(strip=True, separator='\n')
                    if len(text) > 100:  # Только значимые блоки текста
                        text_parts.append(text)
        
        # Если не нашли контентные блоки, берем весь body
        if not text_parts:
            body = soup.find('body')
            if body:
                text = body.get_text(strip=True, separator='\n')
                text_parts.append(text)
        
        # Если все еще нет текста, берем весь документ
        if not text_parts:
            text = soup.get_text(strip=True, separator='\n')
            text_parts.append(text)
        
        # Объединяем текст
        full_text = '\n\n'.join(text_parts)
        
        # Очищаем текст от лишних пробелов и переносов
        import re
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)  # Убираем множественные переносы
        full_text = re.sub(r'\s{2,}', ' ', full_text)     # Убираем множественные пробелы
        
        # Ограничиваем длину текста (опционально)
        if len(full_text) > 100000:
            full_text = full_text[:100000] + "\n\n[Текст обрезан из-за большого объема]"
        
        return full_text.strip()
    
    except requests.exceptions.RequestException as e:
        st.error(f"Ошибка загрузки веб-страницы {url}: {e}")
        return ""
    except Exception as e:
        st.error(f"Ошибка обработки веб-страницы: {e}")
        return ""

def clean_text(text):
    """
    Очищает текст от лишних пробелов и форматирования
    """
    # Убираем множественные переносы строк
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Убираем множественные пробелы
    text = re.sub(r'\s+', ' ', text)
    
    # Убираем спецсимволы, но сохраняем пунктуацию
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    
    # Убираем HTML-сущности
    text = re.sub(r'&[a-z]+;', ' ', text)
    
    # Обрезаем слишком длинные строки
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if len(line) > 1000:
            # Разбиваем слишком длинные строки по предложениям
            sentences = re.split(r'(?<=[.!?])\s+', line)
            cleaned_lines.extend(sentences)
        else:
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    return text.strip()

def is_valid_url(url):
    """
    Проверяет валидность URL
    """
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except:
        return False

def calculate_file_hash(file_path):
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def upload_document(file, file_type, url=None):
    conn = init_db()
    if not conn:
        return False
    
    tmp_file_path = None
    try:
        if url:
            # Для URL создаем временный файл с текстом
            content = extract_text_from_url(url)
            if not content:
                st.error("Не удалось извлечь текст из веб-страницы")
                return False
            
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as tmp_file:
                tmp_file.write(content)
                tmp_file_path = tmp_file.name
                file_hash = hashlib.sha256(url.encode()).hexdigest()
        else:
            # Сохранение файла во временную папку
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
                tmp_file.write(file.read())
                tmp_file_path = tmp_file.name
                file_hash = calculate_file_hash(tmp_file_path)
                content = extract_text_from_file(tmp_file_path, file_type)
        
        # Проверка на дубликаты
        cur = conn.cursor()
        cur.execute("SELECT id FROM documents WHERE file_hash = %s", (file_hash,))
        if cur.fetchone():
            st.warning("Этот документ уже был загружен ранее!")
            return False
        
        # Сохранение в базу данных
        metadata = {
            "source": url if url else file.name,
            "upload_date": datetime.now().isoformat(),
            "original_url": url if url else None
        }
        
        cur.execute("""
            INSERT INTO documents 
            (filename, file_type, file_size, file_hash, content, metadata)
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (
            url if url else file.name,
            "url" if url else file_type,
            len(content.encode('utf-8')) if content else 0,
            file_hash,
            content,
            json.dumps(metadata)
        ))
        
        doc_id = cur.fetchone()[0]
        conn.commit()
        
        if content:
            process_document(content, doc_id, metadata)
            cur.execute("UPDATE documents SET processed = TRUE WHERE id = %s", (doc_id,))
            conn.commit()
        
        st.success("Документ успешно загружен и обработан!")
        return True
        
    except Exception as e:
        st.error(f"Ошибка при загрузке документа: {e}")
        conn.rollback()
        return False
    
    finally:
        conn.close()
        if tmp_file_path and os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)

def process_document(content, doc_id, metadata):
    try:
        doc = Document(
            text=content,
            metadata={
                "doc_id": doc_id,
                **metadata
            }
        )
        
        splitter = SemanticSplitterNodeParser(
            buffer_size=1,
            breakpoint_percentile_threshold=95,
            embed_model=Settings.embed_model
        )
        
        pipeline = IngestionPipeline(
            transformations=[
                splitter,
                Settings.embed_model
            ],
            vector_store=init_vector_store()
        )
        
        nodes = pipeline.run(documents=[doc])
        
        return True
    
    except Exception as e:
        st.error(f"Ошибка обработки документа: {e}")
        return False

def get_documents():
    conn = init_db()
    if not conn:
        return []
    
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT id, filename, file_type, file_size, upload_date, processed
            FROM documents 
            ORDER BY upload_date DESC
        """)
        
        documents = cur.fetchall()
        return documents
    
    except Exception as e:
        st.error(f"Ошибка получения списка документов: {e}")
        return []
    
    finally:
        conn.close()

def delete_document(doc_id):
    conn = init_db()
    if not conn:
        return False
    
    try:
        cur = conn.cursor()
        
        vector_store = init_vector_store()
        if vector_store:
            # Здесь нужно добавить логику удаления чанков по doc_id
            pass
        
        cur.execute("DELETE FROM documents WHERE id = %s", (doc_id,))
        conn.commit()
        
        return True
    
    except Exception as e:
        st.error(f"Ошибка удаления документа: {e}")
        conn.rollback()
        return False
    
    finally:
        conn.close()

@st.cache_resource
def init_rag_system():
    try:
        vector_store = init_vector_store()
        if not vector_store:
            return None
        
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        
        try:
            index = load_index_from_storage(storage_context)
        except:
            index = VectorStoreIndex.from_vector_store(
                vector_store=vector_store,
                embed_model=Settings.embed_model
            )
        
        query_engine = index.as_query_engine(
            llm=Settings.llm,
            similarity_top_k=3,
            system_prompt="""Ты - 'СПИНТехтёнок', цифровой помощник института СПИНТех в НИУ МИЭТ.
            Ты должен отвечать ТОЛЬКО на вопросы по истории Института СПИНТех и о мероприятии 'День СПИНТеха'.
            Если вопрос не связан с этими темами, вежливо откажись отвечать, объяснив свою специализацию.
            Используй предоставленные документы для точных и достоверных ответов."""
        )
        
        return query_engine
    
    except Exception as e:
        st.error(f"Ошибка инициализации RAG системы: {e}")
        return None

def classify_intent(user_input: str, llm: OpenRouter) -> str:
    prompt = f"""Проанализируй запрос пользователя и определи его намерение. Выбери ТОЛЬКО ОДНУ из трех категорий:

1. chitchat - если пользователь:
   - Приветствует или прощается
   - Задает общие вопросы о тебе или твоих возможностях
   - Просто беседует на отвлеченные темы
   - Примеры: "Привет", "Как дела?", "Чем ты занимаешься?", "Расскажи о себе"

2. knowledge - если пользователь:
   - Задает вопросы по истории института СПИНТех
   - Спрашивает о мероприятии "День СПИНТеха"
   - Ищет конкретную информацию о СПИНТехе
   - Примеры: "Когда основан СПИНТех?", "Что будет на Дне СПИНТеха?", "Расскажи о факультетах"

3. out_of_scope - если пользователь:
   - Задает вопросы не по теме СПИНТеха
   - Просит сделать что-то запрещенное или неэтичное
   - Отправляет мусорный текст или бессмыслицу
   - Примеры: "Как приготовить торт?", "Взломай сайт", "asdfghjkl"

Запрос пользователя: "{user_input}"

Твой ответ должен содержать ТОЛЬКО одно слово: chitchat, knowledge или out_of_scope.
Не добавляй никаких пояснений, только категорию."""

    try:
        response = llm.complete(prompt, max_tokens=10)
        intent = response.text.strip().lower()
        
        # Валидация ответа
        valid_intents = ['chitchat', 'knowledge', 'out_of_scope']
        
        # Проверяем, содержит ли ответ одну из валидных категорий
        for valid_intent in valid_intents:
            if valid_intent in intent:
                return valid_intent
        
        # Если ответ не распознан, используем fallback логику
        return fallback_intent_classification(user_input)
        
    except Exception as e:
        st.error(f"Ошибка классификации намерения: {e}")
        raise

def get_intent_specific_response(intent: str, user_input: str, query_engine=None):
    if intent == 'knowledge':
        # Используем RAG для ответов по СПИНТеху
        if query_engine:
            response = query_engine.query(user_input)
            return response.response
        else:
            return "Извините, система поиска информации временно недоступна."
    
    elif intent == 'chitchat':
        # Для беседы используем базовый промпт
        return "Я - СПИНТехтёнок, цифровой помощник института СПИНТех. Помогаю с вопросами об истории института и мероприятии 'День СПИНТеха'. Чем могу помочь?"
    
    elif intent == 'out_of_scope':
        # Вежливый отказ для неподходящих запросов
        return "Извините, я специализируюсь только на вопросах, связанных с институтом СПИНТех и мероприятием 'День СПИНТеха'. Могу ответить только на вопросы по этой теме."

def chat_interface():
    st.title("спинтехтёнок")
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    if "query_engine" not in st.session_state:
        with st.spinner("Инициализация помощника..."):
            st.session_state.query_engine = init_rag_system()
    
    # Инициализация LLM для классификации
    if "intent_llm" not in st.session_state:
        st.session_state.intent_llm = OpenRouter(
            model=Config.LLM_MODEL,
            api_key=Config.OPENROUTER_API_KEY,
            base_url=Config.OPENROUTER_BASE_URL,
            temperature=0.1,
            context_window=4096
        )
    
    # Отображение истории чата
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    if prompt := st.chat_input("Задайте вопрос о СПИНТехе..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("СПИНТехтёнок думает..."):
                try:
                    # 1. Определяем намерение
                    intent = classify_intent(prompt, st.session_state.intent_llm)
                    
                    # 2. Получаем соответствующий ответ
                    answer = get_intent_specific_response(
                        intent, 
                        prompt, 
                        st.session_state.query_engine
                    )
                    
                except Exception as e:
                    answer = f"Произошла ошибка: {str(e)}"
                
                st.markdown(answer)
        
        st.session_state.messages.append({"role": "assistant", "content": answer})

def admin_interface():
    st.title("⚙️ Административная панель")
    
    tab1, tab2, tab3 = st.tabs(["📤 Загрузить документы", "📋 Список документов", "⚙️ Настройки"])
    
    with tab1:
        st.subheader("Загрузка новых документов")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Загрузка файлов")
            file_type = st.selectbox(
                "Тип файла",
                ["pdf", "docx", "txt"]
            )
            
            uploaded_file = st.file_uploader(
                f"Выберите {file_type.upper()} файл",
                type=[file_type]
            )
            
            if uploaded_file and st.button("Загрузить файл"):
                with st.spinner("Обработка файла..."):
                    success = upload_document(uploaded_file, file_type)
                    if success:
                        st.rerun()
        
        with col2:
            st.markdown("#### Загрузка веб-страницы")
            url = st.text_input("URL веб-страницы")
            
            if url and st.button("Загрузить страницу"):
                with st.spinner("Загрузка и обработка страницы..."):
                    success = upload_document(None, "html", url)
                    if success:
                        st.rerun()
        
        st.markdown("---")
        st.info("""
        **Поддерживаемые форматы:**
        - PDF, DOCX, TXT файлы
        - Веб-страницы (через URL)
        
        **Что происходит после загрузки:**
        1. Документ сохраняется в базе данных
        2. Текст извлекается из документа
        3. Текст разбивается на семантические чанки
        4. Чанки сохраняются в векторной базе данных
        """)
    
    with tab2:
        st.subheader("Загруженные документы")
        
        documents = get_documents()
        
        if not documents:
            st.info("Документы не загружены")
        else:
            df = pd.DataFrame(documents, columns=["ID", "Имя файла", "Тип", "Размер", "Дата загрузки", "Обработан"])
            df["Размер"] = df["Размер"].apply(lambda x: f"{x} байт" if x else "N/A")
            df["Дата загрузки"] = pd.to_datetime(df["Дата загрузки"]).dt.strftime("%Y-%m-%d %H:%M")
            df["Статус"] = df["Обработан"].apply(lambda x: "✅" if x else "⏳")
            
            st.dataframe(df.drop(columns=["Обработан"]), width=True)
            
            st.markdown("### Управление документами")
            doc_id_to_delete = st.number_input("ID документа для удаления", min_value=1, step=1)
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("Удалить документ", type="secondary"):
                    if doc_id_to_delete:
                        if delete_document(doc_id_to_delete):
                            st.success("Документ удален!")
                            st.rerun()
                    else:
                        st.warning("Введите ID документа")
            
            with col2:
                if st.button("Обновить список", type="secondary"):
                    st.rerun()
    
    with tab3:
        st.subheader("Настройки системы")
        
        st.markdown("### Конфигурация моделей")
        st.info(f"**LLM модель:** {Config.LLM_MODEL}")
        st.info(f"**Embedding модель:** {Config.EMBEDDING_MODEL}")
        
        st.markdown("### Статистика базы данных")
        conn = init_db()
        if conn:
            cur = conn.cursor()
            
            cur.execute("SELECT COUNT(*) FROM documents")
            doc_count = cur.fetchone()[0]
            
            cur.execute("SELECT COUNT(*) FROM documents WHERE processed = TRUE")
            processed_count = cur.fetchone()[0]
            
            try:
                cur.execute("SELECT COUNT(*) FROM document_chunks")
                chunk_count = cur.fetchone()[0] or 0
            except:
                chunk_count = 0
            
            conn.close()
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Всего документов", doc_count)
            with col2:
                st.metric("Обработано", processed_count)
            with col3:
                st.metric("Чанков в базе", chunk_count)
        
        st.markdown("### Очистка данных")
        if st.button("Очистить историю чата", type="secondary"):
            st.session_state.messages = []
            st.success("История чата очищена!")
        
        if st.button("Переиндексировать документы", type="secondary"):
            st.warning("Эта функция находится в разработке")

def main():
    create_tables()
    
    init_models()
    
    with st.sidebar:
        page = st.radio(
            "Навигация",
            ["💬 Чатбот", "⚙️ Админка"]
        )
    
    if page == "💬 Чатбот":
        chat_interface()
    else:
        admin_interface()

if __name__ == "__main__":
    main()
